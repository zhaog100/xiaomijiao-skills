#!/usr/bin/env python3
"""
长图转Word/PDF工具
支持超长图片自动分块处理
"""

import sys
import os
from pathlib import Path
from PIL import Image

def split_long_image(image_path, max_height=2000):
    """分割超长图片"""
    img = Image.open(image_path)
    width, height = img.size
    
    if height <= max_height:
        return [img]
    
    # 计算分块数量
    chunks = []
    num_chunks = (height + max_height - 1) // max_height
    
    for i in range(num_chunks):
        top = i * max_height
        bottom = min((i + 1) * max_height, height)
        chunk = img.crop((0, top, width, bottom))
        chunks.append(chunk)
    
    return chunks

def to_word(image_path, output_path=None):
    """转换为Word文档"""
    from docx import Document
    from docx.shared import Inches
    
    if output_path is None:
        output_path = Path(image_path).stem + ".docx"
    
    # 分割图片
    chunks = split_long_image(image_path)
    
    # 创建Word文档
    doc = Document()
    doc.add_heading(f'图片文档: {Path(image_path).name}', 0)
    
    for i, chunk in enumerate(chunks, 1):
        if len(chunks) > 1:
            doc.add_heading(f'分块 {i}/{len(chunks)}', level=1)
        
        # 保存临时图片
        temp_path = f"/tmp/chunk_{i}.png"
        chunk.save(temp_path)
        
        # 添加到文档（宽度6英寸，高度自动）
        doc.add_picture(temp_path, width=Inches(6))
        doc.add_paragraph()  # 空行
        
        # 删除临时文件
        os.remove(temp_path)
    
    doc.save(output_path)
    print(f"✅ Word文档已生成: {output_path}")
    return output_path

def to_pdf(image_path, output_path=None):
    """转换为PDF文档"""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Image as RLImage, PageBreak
    from reportlab.lib.units import inch
    
    if output_path is None:
        output_path = Path(image_path).stem + ".pdf"
    
    # 分割图片
    chunks = split_long_image(image_path)
    
    # 创建PDF
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=0.5*inch,
        leftMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    story = []
    temp_files = []
    for i, chunk in enumerate(chunks):
        if i > 0:
            story.append(PageBreak())
        
        # 保存临时图片
        temp_path = f"/tmp/chunk_{i}.png"
        chunk.save(temp_path)
        temp_files.append(temp_path)
        
        # 添加图片（适应页面尺寸）
        img = RLImage(temp_path)
        
        # 计算可用空间
        available_width = 7*inch
        available_height = 9*inch  # A4高度减去边距
        
        # 根据图片比例计算缩放
        aspect = chunk.size[1] / chunk.size[0]
        img.drawWidth = available_width
        img.drawHeight = img.drawWidth * aspect
        
        # 如果高度超出，按高度缩放
        if img.drawHeight > available_height:
            img.drawHeight = available_height
            img.drawWidth = img.drawHeight / aspect
        
        story.append(img)
    
    doc.build(story)
    
    # 删除临时文件
    for temp_path in temp_files:
        os.remove(temp_path)
    print(f"✅ PDF文档已生成: {output_path}")
    return output_path

def main():
    if len(sys.argv) < 3:
        print("用法: image-to-doc.py <图片路径> <word|pdf|both> [输出路径]")
        print("示例: image-to-doc.py screenshot.png both")
        sys.exit(1)
    
    image_path = sys.argv[1]
    format_type = sys.argv[2].lower()
    output_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    if not os.path.exists(image_path):
        print(f"❌ 图片不存在: {image_path}")
        sys.exit(1)
    
    if format_type in ['word', 'both']:
        to_word(image_path, output_path.replace('.pdf', '.docx') if output_path else None)
    
    if format_type in ['pdf', 'both']:
        to_pdf(image_path, output_path.replace('.docx', '.pdf') if output_path else None)

if __name__ == "__main__":
    main()
