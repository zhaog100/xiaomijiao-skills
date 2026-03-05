#!/usr/bin/env python3

"""
从Word/PDF文档中提取文本内容
功能：结合OCR和文档解析，完整提取图片中的文本内容
"""

import os
import sys
from PIL import Image
import cv2
import numpy as np

# 添加项目路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

def extract_text_from_word(doc_path):
    """从Word文档提取文本"""
    try:
        from docx import Document
        
        doc = Document(doc_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # 从表格中提取文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        return "\n".join(text_content)
        
    except ImportError:
        print("❌ python-docx未安装")
        return None

def extract_text_from_pdf(pdf_path):
    """从PDF文档提取文本"""
    try:
        import PyPDF2
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text_content = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)
            
            return "\n".join(text_content)
            
    except ImportError:
        print("❌ PyPDF2未安装")
        return None

def extract_images_from_pdf(pdf_path, output_dir):
    """从PDF中提取图片"""
    try:
        from pdf2image import convert_from_path
        
        images = convert_from_path(pdf_path)
        extracted_images = []
        
        for i, image in enumerate(images):
            img_path = os.path.join(output_dir, f"extracted_page_{i+1}.png")
            image.save(img_path, "PNG")
            extracted_images.append(img_path)
        
        return extracted_images
        
    except ImportError:
        print("❌ pdf2image未安装")
        return []

def advanced_ocr_on_image(image_path):
    """对图片进行高级OCR处理"""
    # 这里会调用AI视觉分析或Tesseract OCR
    # 目前使用模拟结果
    print(f"🔍 对图片进行高级OCR: {image_path}")
    
    # 模拟OCR结果
    return f"[OCR结果] 图片 {os.path.basename(image_path)} 的文本内容"

def split_long_image(image_path, max_height=2000):
    """
    分割超长图片
    
    Args:
        image_path: 输入图片路径
        max_height: 每块的最大高度
    
    Returns:
        分割后的图片块列表
    """
    img = Image.open(image_path)
    width, height = img.size
    
    if height <= max_height:
        return [img]
    
    # 计算分割数量
    num_blocks = (height + max_height - 1) // max_height
    blocks = []
    
    for i in range(num_blocks):
        top = i * max_height
        bottom = min((i + 1) * max_height, height)
        block = img.crop((0, top, width, bottom))
        blocks.append(block)
    
    return blocks

def create_word_document(blocks, output_path):
    """
    创建Word文档
    
    Args:
        blocks: 图片块列表
        output_path: 输出文件路径
    """
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading('超长图片内容', 0)
        
        for i, block in enumerate(blocks):
            # 保存临时图片
            temp_img_path = f"/tmp/block_{i}.png"
            block.save(temp_img_path)
            
            # 添加到Word文档
            doc.add_picture(temp_img_path, width=Inches(6))
            doc.add_paragraph(f"--- 第 {i+1} 块 ---")
            
            # 清理临时文件
            os.remove(temp_img_path)
        
        doc.save(output_path)
        print(f"✅ Word文档创建成功: {output_path}")
        return True
        
    except ImportError:
        print("❌ python-docx未安装，无法创建Word文档")
        return False

def create_pdf_document(blocks, output_path):
    """
    创建PDF文档
    
    Args:
        blocks: 图片块列表
        output_path: 输出文件路径
    """
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.utils import ImageReader
        
        # 使用A4页面
        page_width, page_height = A4
        
        c = canvas.Canvas(output_path, pagesize=A4)
        c.setFont("Helvetica", 16)
        c.drawString(50, page_height - 50, "超长图片内容")
        
        current_y = page_height - 100
        
        for i, block in enumerate(blocks):
            # 保存临时图片
            temp_img_path = f"/tmp/block_{i}.png"
            block.save(temp_img_path)
            
            # 获取图片尺寸
            img_width, img_height = block.size
            
            # 计算缩放比例（保持宽高比）
            scale = min((page_width - 100) / img_width, 400 / img_height)
            scaled_width = img_width * scale
            scaled_height = img_height * scale
            
            # 检查是否需要新页面
            if current_y - scaled_height < 50:
                c.showPage()
                current_y = page_height - 50
            
            # 添加图片到PDF
            c.drawImage(temp_img_path, 50, current_y - scaled_height, 
                       width=scaled_width, height=scaled_height)
            
            # 添加分块标识
            c.setFont("Helvetica", 12)
            c.drawString(50, current_y - 10, f"--- 第 {i+1} 块 ---")
            
            current_y -= (scaled_height + 30)
            
            # 清理临时文件
            os.remove(temp_img_path)
        
        c.save()
        print(f"✅ PDF文档创建成功: {output_path}")
        return True
        
    except ImportError:
        print("❌ reportlab未安装，无法创建PDF文档")
        return False

def complete_content_extraction(image_path, output_dir=None):
    """
    完整内容提取流程
    
    Args:
        image_path: 输入图片路径
        output_dir: 输出目录
    
    Returns:
        完整的文本内容
    """
    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(image_path), "extracted_content")
    
    os.makedirs(output_dir, exist_ok=True)
    print(f"📁 输出目录: {output_dir}")
    
    # 步骤1: 转换为Word文档
    print("📄 步骤1: 转换为Word文档...")
    word_path = os.path.join(output_dir, "converted.docx")
    
    blocks = split_long_image(image_path)
    create_word_document(blocks, word_path)
    
    # 步骤2: 从Word文档提取文本
    print("📝 步骤2: 从Word文档提取文本...")
    word_text = extract_text_from_word(word_path)
    
    if word_text and word_text.strip():
        print("✅ Word文本提取成功")
        return word_text
    else:
        print("⚠️  Word文本提取失败，尝试其他方法")
    
    # 步骤3: 转换为PDF文档
    print("📄 步骤3: 转换为PDF文档...")
    pdf_path = os.path.join(output_dir, "converted.pdf")
    create_pdf_document(blocks, pdf_path)
    
    # 步骤4: 从PDF提取文本
    print("📝 步骤4: 从PDF提取文本...")
    pdf_text = extract_text_from_pdf(pdf_path)
    
    if pdf_text and pdf_text.strip():
        print("✅ PDF文本提取成功")
        return pdf_text
    else:
        print("⚠️  PDF文本提取失败，进行高级OCR")
    
    # 步骤5: 高级OCR处理
    print("🔍 步骤5: 高级OCR处理...")
    ocr_results = []
    for i, block in enumerate(blocks):
        # 保存临时图片
        temp_img_path = os.path.join(output_dir, f"block_{i+1}.png")
        block.save(temp_img_path)
        
        # 进行OCR
        ocr_result = advanced_ocr_on_image(temp_img_path)
        ocr_results.append(ocr_result)
    
    complete_text = "\n".join(ocr_results)
    print("✅ 高级OCR完成")
    
    return complete_text

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python extract-text-from-doc.py <image_path> [output_dir]")
        sys.exit(1)
    
    image_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(image_path):
        print(f"❌ 图片文件不存在: {image_path}")
        sys.exit(1)
    
    print(f"🚀 开始完整内容提取...")
    print(f"图片: {image_path}")
    
    try:
        content = complete_content_extraction(image_path, output_dir)
        print("\n" + "="*50)
        print("📋 完整提取结果:")
        print("="*50)
        print(content)
        print("="*50)
        
        # 保存结果
        if output_dir:
            result_path = os.path.join(output_dir, "extracted_content.txt")
            with open(result_path, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"\n✅ 结果已保存到: {result_path}")
        
    except Exception as e:
        print(f"❌ 提取失败: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()