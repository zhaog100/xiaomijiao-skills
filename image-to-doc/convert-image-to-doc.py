#!/usr/bin/env python3

"""
图片转文档脚本
功能：将超长图片转换为Word或PDF格式，便于内容读取
"""

import os
import sys
from PIL import Image
import cv2
import numpy as np

def split_long_image(image_path, max_height=2000):
    """
    分割超长图片
    
    Args:
        image_path: 输入图片路径
        max_height: 每块的最大高度
    
    Returns:
        分割后的图片块列表
    """
    img = Image.open(image_path)
    width, height = img.size
    
    if height <= max_height:
        return [img]
    
    # 计算分割数量
    num_blocks = (height + max_height - 1) // max_height
    blocks = []
    
    for i in range(num_blocks):
        top = i * max_height
        bottom = min((i + 1) * max_height, height)
        block = img.crop((0, top, width, bottom))
        blocks.append(block)
    
    return blocks

def create_word_document(blocks, output_path):
    """
    创建Word文档
    
    Args:
        blocks: 图片块列表
        output_path: 输出文件路径
    """
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading('超长图片内容', 0)
        
        for i, block in enumerate(blocks):
            # 保存临时图片
            temp_img_path = f"/tmp/block_{i}.png"
            block.save(temp_img_path)
            
            # 添加到Word文档
            doc.add_picture(temp_img_path, width=Inches(6))
            doc.add_paragraph(f"--- 第 {i+1} 块 ---")
            
            # 清理临时文件
            os.remove(temp_img_path)
        
        doc.save(output_path)
        print(f"✅ Word文档创建成功: {output_path}")
        return True
        
    except ImportError:
        print("❌ python-docx未安装，无法创建Word文档")
        return False

def create_pdf_document(blocks, output_path):
    """
    创建PDF文档
    
    Args:
        blocks: 图片块列表
        output_path: 输出文件路径
    """
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.utils import ImageReader
        
        # 使用A4页面
        page_width, page_height = A4
        
        c = canvas.Canvas(output_path, pagesize=A4)
        c.setFont("Helvetica", 16)
        c.drawString(50, page_height - 50, "超长图片内容")
        
        current_y = page_height - 100
        
        for i, block in enumerate(blocks):
            # 保存临时图片
            temp_img_path = f"/tmp/block_{i}.png"
            block.save(temp_img_path)
            
            # 获取图片尺寸
            img_width, img_height = block.size
            
            # 计算缩放比例（保持宽高比）
            scale = min((page_width - 100) / img_width, 400 / img_height)
            scaled_width = img_width * scale
            scaled_height = img_height * scale
            
            # 检查是否需要新页面
            if current_y - scaled_height < 50:
                c.showPage()
                current_y = page_height - 50
            
            # 添加图片到PDF
            c.drawImage(temp_img_path, 50, current_y - scaled_height, 
                       width=scaled_width, height=scaled_height)
            
            # 添加分块标识
            c.setFont("Helvetica", 12)
            c.drawString(50, current_y - 10, f"--- 第 {i+1} 块 ---")
            
            current_y -= (scaled_height + 30)
            
            # 清理临时文件
            os.remove(temp_img_path)
        
        c.save()
        print(f"✅ PDF文档创建成功: {output_path}")
        return True
        
    except ImportError:
        print("❌ reportlab未安装，无法创建PDF文档")
        return False

def main():
    """主函数"""
    if len(sys.argv) < 3:
        print("用法: python convert-image-to-doc.py <image_path> <output_path> [format]")
        print("格式: word (默认) 或 pdf")
        sys.exit(1)
    
    image_path = sys.argv[1]
    output_path = sys.argv[2]
    format_type = sys.argv[3] if len(sys.argv) > 3 else "word"
    
    if not os.path.exists(image_path):
        print(f"❌ 图片文件不存在: {image_path}")
        sys.exit(1)
    
    print(f"🚀 开始转换图片到{format_type.upper()}...")
    print(f"图片: {image_path}")
    print(f"输出: {output_path}")
    
    # 分割图片
    blocks = split_long_image(image_path)
    print(f"✂️  分割成 {len(blocks)} 块")
    
    # 创建文档
    if format_type.lower() == "pdf":
        success = create_pdf_document(blocks, output_path)
    else:
        success = create_word_document(blocks, output_path)
    
    if success:
        print("🎉 转换完成！")
    else:
        print("❌ 转换失败！")

if __name__ == "__main__":
    main()